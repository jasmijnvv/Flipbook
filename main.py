import textract
import regex as re

import queue
from spire.doc import *
from spire.doc.common import *

from pathlib import Path

from docxcompose.composer import Composer
import docx

import math
#math.ceil maakt er een mooi rond getal van (voor het delen van het aantal karakters door 255 dat waarschijnlijk geen mooi getal oplevert)

import HTML 

def merge_text(splitted_text):
    text = ""
    for word in splitted_text:
        text += word + " "
    return text

def create_html_file(file_path, titels, page_texts, image_paths, namen):
    text = ""
    text += HTML.doc_begin

    text += HTML.create_page_cover("TITEL", "image")
    for titel, page_text, image_path, naam in zip (titels, page_texts, image_paths, namen):
        split_text = page_text.split(" ")
        words_per_page = 50
        if len(split_text) > words_per_page:
            aantal_paginas = math.ceil(len(split_text)/words_per_page)
            for pagina in range(aantal_paginas):
                page_text_short = split_text[pagina*words_per_page: (pagina+1)*words_per_page]
                page_text_short = merge_text(page_text_short)
                if pagina == 0:
                    text += HTML.create_page(titel, page_text_short, "", "")
                elif pagina < aantal_paginas -1: 
                    text += HTML.create_page("", page_text_short, "", "")
                elif pagina == aantal_paginas -1:
                    text += HTML.create_page("", page_text_short, image_path, naam)
            continue
        text += HTML.create_page(titel, page_text, image_path, naam)

    text += HTML.create_page("Einde", "", "", "")

    text += HTML.doc_end

    with open(file_path, "w") as file:
        file.write(text)

#Code from https://www.freecodecamp.org/news/merge-word-documents-in-python/ 
def merge_docs_with_page_breaks(output_path, *files_list):
    files_list = files_list[0]
    base_doc = docx.Document(files_list[0])
    composer = Composer(base_doc)

    for file in files_list[1:]:
        doc = docx.Document(file)

        # adding page break before merging each document
        base_doc.add_page_break()
        composer.append(doc)

    composer.save(output_path)
    

# Code from https://www.e-iceblue.com/Tutorials/Python/Spire.Doc-for-Python/Program-Guide/Text/Python-Extract-Text-and-Images-from-Word-Documents.html
def extract_images(file_name: Path):
        # Create a Document object
    doc = Document()

    # Load a Word file
    doc.LoadFromFile(str(file_name))

    # Create a Queue object
    nodes = queue.Queue()
    nodes.put(doc)

    # Create a list
    images = []

    while nodes.qsize() > 0:
        node = nodes.get()

        # Loop through the child objects in the document
        for i in range(node.ChildObjects.Count):
            child = node.ChildObjects.get_Item(i)

            # Determine if a child object is a picture
            if child.DocumentObjectType == DocumentObjectType.Picture:
                picture = child if isinstance(child, DocPicture) else None
                dataBytes = picture.ImageBytes

                # Add the image data to the list 
                images.append(dataBytes)
            
            elif isinstance(child, ICompositeObject):
                nodes.put(child if isinstance(child, ICompositeObject) else None)

    # Loop through the images in the list
    image_paths = []
    for i, item in enumerate(images):
        fileName = f"{file_name.stem}_image-{i}.png"
        image_paths.append(fileName)
        with open(fileName,'wb') as imageFile:

            # Write the image to a specified path
            imageFile.write(item)
    doc.Close()
    return image_paths

files_list = []
directory = Path(r"C:\Users\Jasmijn\Documents\Dikkiedik")
for file_path in directory.iterdir():
     if file_path.is_file():
        files_list.append(file_path.resolve())
        #print(files_list)


file_name = r"output_path.docx"
merge_docs_with_page_breaks(file_name, files_list)

titels = []
teksten = []
namen = []
images = []
for file in files_list:
    text = textract.process(filename=file)

    text = text.decode('utf-8')

    titel = re.findall(r'(?<=Titel:)[^.]*(?=\n)', text)
    tekst = re.findall(r'(?<=Tekst:)[\s\S]*?(?=Foto)', text)
    naam = re.findall(r'(?<=Naam:)[\s\S]*?(?=Klaar!)',text)

    #print(f"TITEL: {titels.group(0)}")
    #print(f"TEKST: {teksten}")
    #print(f"NAAM: {namen}")
    image = extract_images(file)

    #for text in teksten:
        #print(f"{len(text)}")


    titels.append(titel[0])
    teksten.append(tekst[0])
    namen.append(naam[0])
    if len(image) > 0:
        images.append(image[0])
    else:
        images.append("")


html_file_path = "book.html"
create_html_file(html_file_path, titels, teksten, images, namen)

